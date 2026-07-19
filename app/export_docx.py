"""Geração do prontuário em documento Word (.docx)."""

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor

from . import config

CAMPOS = [
    ("queixa_principal", "Queixa principal"),
    ("historia_doenca_atual", "História da doença atual"),
    ("antecedentes", "Antecedentes"),
    ("exame_fisico", "Exame físico"),
    ("hipoteses_diagnosticas", "Hipóteses diagnósticas"),
    ("conduta", "Conduta"),
    ("observacoes", "Observações"),
]


def _formatar_data(iso):
    try:
        return datetime.fromisoformat(iso).strftime("%d/%m/%Y %H:%M")
    except (ValueError, TypeError):
        return iso or ""


def gerar(consulta):
    """Devolve um BytesIO com o .docx do prontuário da consulta."""
    cfg = config.carregar()
    doc = Document()

    clinica = (cfg.get("clinica") or "").strip()
    if clinica:
        cab = doc.add_paragraph()
        run = cab.add_run(clinica)
        run.bold = True
        run.font.size = Pt(13)

    titulo = doc.add_heading("Prontuário", level=0)
    titulo.runs[0].font.color.rgb = RGBColor(0x11, 0x8A, 0x52)

    # Metadados
    meta = doc.add_paragraph()
    meta.add_run("Paciente: ").bold = True
    meta.add_run(consulta.get("paciente") or "Não informado")
    if consulta.get("especialidade"):
        meta.add_run("\nEspecialidade: ").bold = True
        meta.add_run(consulta["especialidade"])
    meta.add_run("\nData: ").bold = True
    meta.add_run(_formatar_data(consulta.get("criado_em")))

    doc.add_paragraph()

    pront = consulta.get("prontuario") or {}
    for chave, rotulo in CAMPOS:
        doc.add_heading(rotulo, level=2)
        valor = pront.get(chave)
        if chave == "hipoteses_diagnosticas":
            itens = valor if isinstance(valor, list) else []
            if itens:
                for h in itens:
                    doc.add_paragraph(str(h), style="List Bullet")
            else:
                doc.add_paragraph("Não informado")
        else:
            doc.add_paragraph(str(valor) if valor else "Não informado")

    rodape = doc.add_paragraph()
    nota = rodape.add_run(
        "\nDocumento gerado como apoio à documentação clínica. A revisão e a "
        "validação são de responsabilidade do profissional de saúde."
    )
    nota.italic = True
    nota.font.size = Pt(8)
    nota.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
